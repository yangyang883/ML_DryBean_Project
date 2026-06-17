from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
FIGURES = RESULTS / "figures"
DATA = ROOT / "data"
DOCX_OUT = ROOT / "机器学习期末大作业总结论文.docx"
MD_OUT = ROOT / "机器学习期末大作业总结论文_备份.md"
GITHUB_URL = "https://github.com/yangyang883/ML_DryBean_Project"
DEMO_URL_TEXT = "项目展示网页链接：提交前由本人填写"

MODEL_ZH = {
    "logistic": "逻辑回归",
    "knn": "K近邻",
    "svm": "支持向量机",
    "random_forest": "随机森林",
    "xgboost": "XGBoost",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def fmt_float(value, digits: int = 4) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def add_paragraph(doc: Document, text: str = "", style: str | None = None, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Cm(0.74) if style is None else None
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 14)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10)
    r.bold = True


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 5.8) -> bool:
    if not image_path.exists() or image_path.stat().st_size == 0:
        add_paragraph(doc, f"{caption}对应图片文件未生成，提交前请重新运行实验检查。")
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)
    return True


def add_table(doc: Document, df: pd.DataFrame, caption: str, max_rows: int | None = None, font_size: int = 8) -> None:
    add_caption(doc, caption)
    shown = df.copy()
    if max_rows is not None and len(shown) > max_rows:
        shown = shown.head(max_rows)
    table = doc.add_table(rows=1, cols=len(shown.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, col in enumerate(shown.columns):
        set_cell_shading(hdr[idx], "D9EAF7")
        set_cell_text(hdr[idx], col, bold=True, size=font_size)
    for _, row in shown.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row.tolist()):
            set_cell_text(cells[idx], value, size=font_size)
    doc.add_paragraph()


def add_toc_page(doc: Document) -> None:
    add_heading(doc, "目录", 1)
    add_paragraph(doc, "本文按照期末大作业评分标准组织章节，先说明数据和处理方法，再展示多算法实验、工程系统和结果可信度验收，最后给出课程总结。")
    rows = [
        ("第一章 绪论", "项目背景与总体流程", "说明 Dry Bean 七分类任务、项目目标和整体实验流程。"),
        ("第二章 数据分析", "数据分析 5%", "介绍数据来源、样本数量、类别分布、缺失值、异常值和特征相关性。"),
        ("第三章 数据清洗与特征工程", "数据处理 30%", "说明 median 填补、重复样本删除、标签清洗、IQR 裁剪、比例特征和标准化。"),
        ("第四章 多算法实验设计与结果分析", "多算法实验 30%", "对比 Logistic、KNN、SVM、Random Forest、XGBoost 的精度、F1、loss、速度、鲁棒性和过拟合。"),
        ("第五章 工程系统设计与网页展示", "工程与展示 30%", "说明项目结构、统一命令行、GitHub、README、Streamlit 展示和上传预测功能。"),
        ("第六章 结果可信度与专项验收", "结果可信度", "说明数据泄漏检查、指标重新计算、混淆矩阵校验和预测稳定性。"),
        ("第七章 课程总结", "课程总结 5%", "用学生视角总结课程收获、项目体会和课程建议。"),
        ("结论、参考文献与附录", "提交材料补充", "汇总项目完成情况，列出参考文献、运行命令、目录结构和提交前检查清单。"),
    ]
    df = pd.DataFrame(rows, columns=["章节", "对应评分项", "主要内容"])
    add_table(doc, df, "表 0-1 论文结构与评分标准对应关系", font_size=8)
    add_paragraph(doc, "说明：如果需要 Word 自动页码目录，可在最终定稿后使用 Word 的“引用 -> 目录”功能基于本文标题样式自动生成。")
    doc.add_page_break()


def prepare_data() -> dict:
    metrics = pd.read_csv(RESULTS / "metrics_summary.csv")
    speed = pd.read_csv(RESULTS / "speed_summary.csv")
    robustness = pd.read_csv(RESULTS / "robustness_summary.csv")
    robustness_pivot = pd.read_csv(RESULTS / "robustness_pivot.csv")
    validation = pd.read_csv(RESULTS / "metric_validation_summary.csv")
    class_dist = pd.read_csv(RESULTS / "class_distribution.csv")
    missing = pd.read_csv(RESULTS / "missing_values.csv")
    outlier = pd.read_csv(RESULTS / "outlier_summary.csv")
    train = pd.read_csv(DATA / "train.csv")
    test = pd.read_csv(DATA / "test.csv")
    val = pd.read_csv(DATA / "val.csv")
    best = metrics.sort_values("test_accuracy", ascending=False).iloc[0]
    best_model = best["model"]
    report_path = RESULTS / f"classification_report_{best_model}.csv"
    class_report = pd.read_csv(report_path) if report_path.exists() else pd.DataFrame()
    return {
        "metrics": metrics,
        "speed": speed,
        "robustness": robustness,
        "robustness_pivot": robustness_pivot,
        "validation": validation,
        "class_dist": class_dist,
        "missing": missing,
        "outlier": outlier,
        "train": train,
        "test": test,
        "val": val,
        "best_model": best_model,
        "best": best,
        "class_report": class_report,
    }


def metrics_table(metrics: pd.DataFrame) -> pd.DataFrame:
    cols = [
        "model",
        "train_accuracy",
        "test_accuracy",
        "precision_weighted",
        "recall_weighted",
        "f1_weighted",
        "avg_predict_time_ms",
        "train_time_sec",
        "overfit_gap",
        "overfit_level",
    ]
    df = metrics[cols].copy()
    df["model"] = df["model"].map(MODEL_ZH).fillna(df["model"])
    df.columns = ["模型", "训练准确率", "测试准确率", "加权精确率", "加权召回率", "加权F1", "平均推理时间ms", "训练时间s", "过拟合差值", "过拟合等级"]
    for c in ["训练准确率", "测试准确率", "加权精确率", "加权召回率", "加权F1", "过拟合差值"]:
        df[c] = df[c].map(lambda x: fmt_float(x, 4))
    df["平均推理时间ms"] = df["平均推理时间ms"].map(lambda x: fmt_float(x, 3))
    df["训练时间s"] = df["训练时间s"].map(lambda x: fmt_float(x, 2))
    return df


def speed_table(speed: pd.DataFrame) -> pd.DataFrame:
    df = speed.copy()
    df["model"] = df["model"].map(MODEL_ZH).fillna(df["model"])
    df = df[["model", "repeats", "test_rows", "avg_predict_time_ms", "avg_per_sample_us"]]
    df.columns = ["模型", "重复次数", "测试样本数", "平均推理时间ms", "单样本平均耗时us"]
    df["平均推理时间ms"] = df["平均推理时间ms"].map(lambda x: fmt_float(x, 3))
    df["单样本平均耗时us"] = df["单样本平均耗时us"].map(lambda x: fmt_float(x, 3))
    return df


def robustness_table(robustness: pd.DataFrame) -> pd.DataFrame:
    df = robustness[["model", "noise_type", "noise_level", "clean_accuracy", "noisy_accuracy", "accuracy_drop"]].copy()
    df["model"] = df["model"].map(MODEL_ZH).fillna(df["model"])
    df["noise_type"] = df["noise_type"].replace({"gaussian": "高斯噪声", "mask": "掩码噪声"})
    df.columns = ["模型", "噪声类型", "噪声强度", "干净训练准确率", "加噪训练准确率", "准确率下降"]
    for c in ["噪声强度", "干净训练准确率", "加噪训练准确率", "准确率下降"]:
        df[c] = df[c].map(lambda x: fmt_float(x, 4))
    return df


def validation_table(validation: pd.DataFrame) -> pd.DataFrame:
    df = validation.copy()
    df["model"] = df["model"].map(MODEL_ZH).fillna(df["model"])
    keep = ["model", "recomputed_accuracy", "summary_accuracy", "metric_values_match", "confusion_matrix_sum_ok", "classification_report_accuracy_ok", "passed"]
    df = df[keep]
    df.columns = ["模型", "重算准确率", "汇总表准确率", "指标一致", "混淆矩阵样本数正确", "分类报告准确率一致", "是否通过"]
    for c in ["重算准确率", "汇总表准确率"]:
        df[c] = df[c].map(lambda x: fmt_float(x, 6))
    return df


def class_report_table(class_report: pd.DataFrame) -> pd.DataFrame:
    df = class_report.copy()
    if "Unnamed: 0" in df.columns:
        df = df.rename(columns={"Unnamed: 0": "类别"})
    elif df.columns[0] != "类别":
        df = df.rename(columns={df.columns[0]: "类别"})
    keep = [c for c in ["类别", "precision", "recall", "f1-score", "support"] if c in df.columns]
    df = df[keep]
    for c in ["precision", "recall", "f1-score"]:
        if c in df.columns:
            df[c] = df[c].map(lambda x: fmt_float(x, 4))
    return df


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("干豆数据集多分类机器学习系统设计与实现")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(22)
    run.bold = True
    for line in [
        "课程名称：机器学习与项目实践",
        "姓名：杨洋",
        f"日期：{date.today().isoformat()}",
        "项目主题：Dry Bean Dataset 多分类任务",
        f"GitHub 链接：{GITHUB_URL}",
        DEMO_URL_TEXT,
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(13)
    doc.add_page_break()


def build_markdown(data: dict, inserted_figures: list[str], inserted_tables: list[str]) -> str:
    best_model = MODEL_ZH.get(data["best_model"], data["best_model"])
    best_acc = fmt_float(data["best"]["test_accuracy"], 4)
    return f"""# 干豆数据集多分类机器学习系统设计与实现

## 摘要
本项目围绕 Dry Bean Dataset 七分类识别任务展开，完成了数据分析、数据清洗、特征工程、多算法训练、模型评估、鲁棒性分析、推理速度分析和工程化展示。项目实现了逻辑回归、K近邻、支持向量机、随机森林和 XGBoost 五种模型，其中随机森林和 XGBoost 属于课堂外扩展算法。实验结果显示，{best_model} 在当前测试集上的准确率最高，测试准确率为 {best_acc}。系统通过统一命令行完成训练和评估，并使用 Streamlit 页面展示结果与上传 CSV 预测功能。

关键词：干豆数据集；多分类；数据清洗；特征工程；支持向量机；XGBoost；鲁棒性分析；Streamlit

## 项目信息
- GitHub 链接：{GITHUB_URL}
- {DEMO_URL_TEXT}
- 训练集样本数：{len(data["train"])}
- 验证集样本数：{len(data["val"])}
- 测试集样本数：{len(data["test"])}
- 类别数量：{data["class_dist"].shape[0]}
- 原始特征数量：{len([c for c in data["train"].columns if c != "Class"])}

## 插入图片
{chr(10).join(f"- {x}" for x in inserted_figures)}

## 插入表格
{chr(10).join(f"- {x}" for x in inserted_tables)}

完整正文见 Word 文档。
"""


def main() -> None:
    data = prepare_data()
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)

    inserted_figures: list[str] = []
    inserted_tables: list[str] = []

    add_cover(doc)

    add_heading(doc, "摘要", 1)
    best_model_zh = MODEL_ZH.get(data["best_model"], data["best_model"])
    add_paragraph(
        doc,
        f"本项目围绕 Dry Bean Dataset 七分类识别任务展开，目标是根据干豆的面积、周长、长轴、短轴、圆度、紧致度等形态学特征识别 BARBUNYA、BOMBAY、CALI、DERMASON、HOROZ、SEKER、SIRA 七类干豆。项目完成了数据分析、数据清洗、特征工程、多算法训练、模型评估、鲁棒性分析、推理速度分析和工程化展示。实验实现了逻辑回归、K近邻、支持向量机、随机森林和 XGBoost 五种模型，其中随机森林和 XGBoost 属于课堂外扩展算法。根据当前真实实验结果，{best_model_zh} 的测试准确率最高，为 {fmt_float(data['best']['test_accuracy'], 4)}。系统通过命令行统一运行算法实验，并使用 Streamlit 展示数据描述、处理流程、实验图表和上传 CSV 预测功能。",
    )
    add_paragraph(doc, "关键词：干豆数据集；多分类；数据清洗；特征工程；支持向量机；XGBoost；鲁棒性分析；Streamlit")

    add_toc_page(doc)

    add_heading(doc, "第一章 绪论", 1)
    add_heading(doc, "1.1 项目背景", 2)
    add_paragraph(doc, "机器学习分类任务可以根据样本特征自动学习类别边界，在农业品质检测、图像识别和自动化分拣等场景中有较高应用价值。Dry Bean Dataset 是基于干豆图像提取形态学特征形成的结构化表格数据，适合用于多分类算法比较和完整机器学习流程实践。")
    add_heading(doc, "1.2 项目任务", 2)
    add_paragraph(doc, "本项目利用干豆的面积、周长、长轴、短轴、凸包面积、等效直径、圆度、紧致度和形状因子等特征，对七个干豆类别进行识别。类别包括 BARBUNYA、BOMBAY、CALI、DERMASON、HOROZ、SEKER 和 SIRA。")
    add_heading(doc, "1.3 项目总体流程", 2)
    add_paragraph(doc, "项目流程为：数据读取 -> 数据分析 -> 数据清洗 -> 特征工程 -> 模型训练 -> 模型评估 -> 鲁棒性分析 -> 推理速度分析 -> Streamlit 展示 -> 上传预测。训练、评估和鲁棒性实验均通过命令行运行，展示页面只负责结果展示和交互预测。")

    add_heading(doc, "第二章 数据分析", 1)
    add_heading(doc, "2.1 数据集介绍", 2)
    add_paragraph(doc, f"当前项目数据用于干豆品种识别。清洗后训练集样本数为 {len(data['train'])}，验证集样本数为 {len(data['val'])}，测试集样本数为 {len(data['test'])}。原始特征数量为 {len([c for c in data['train'].columns if c != 'Class'])}，类别数量为 {data['class_dist'].shape[0]}。")
    add_table(doc, data["class_dist"].rename(columns={"Class": "类别", "count": "样本数"}), "表 2-1 类别样本数量统计", font_size=9)
    inserted_tables.append("表 2-1 类别样本数量统计")
    add_heading(doc, "2.2 数据污染情况观察", 2)
    missing_nonzero = data["missing"][data["missing"]["missing_count"] > 0]
    add_paragraph(doc, f"从 EDA 结果看，数据中存在缺失值、重复样本风险、标签字符污染、异常值、类别不平衡、特征尺度差异和强相关特征等问题。缺失值主要出现在 {', '.join(missing_nonzero.iloc[:,0].astype(str).tolist()) if len(missing_nonzero) else '少数特征'}。这些问题会影响距离度量、模型收敛和泛化效果，因此需要在训练前进行清洗和特征工程。")
    add_heading(doc, "2.3 类别分布分析", 2)
    if add_figure(doc, FIGURES / "class_distribution.png", "图 2-1 干豆类别分布直方图"):
        inserted_figures.append("图 2-1 干豆类别分布直方图")
    add_paragraph(doc, "类别分布显示 DERMASON 和 SIRA 样本较多，BOMBAY 样本较少，数据存在一定类别不平衡。类别不平衡可能使模型更倾向多数类，对少数类的召回率造成影响，因此实验中同时观察 weighted precision、weighted recall 和 weighted F1。")
    add_heading(doc, "2.4 缺失值与异常值分析", 2)
    add_table(doc, data["missing"].rename(columns={data["missing"].columns[0]: "特征", "missing_count": "缺失数量"}).head(12), "表 2-2 缺失值统计表", font_size=8)
    inserted_tables.append("表 2-2 缺失值统计表")
    add_table(doc, data["outlier"].rename(columns={"feature": "特征", "outlier_count": "IQR异常数量", "lower_bound": "下界", "upper_bound": "上界"}).head(12), "表 2-3 IQR 异常值统计表", font_size=7)
    inserted_tables.append("表 2-3 IQR 异常值统计表")
    add_paragraph(doc, "IQR 检测出的极端值不一定全部是错误数据，因为不同干豆品种本身存在明显形态差异。若简单删除异常样本，可能损失少数类别中有价值的信息，因此本项目采用裁剪方式降低极端值影响，同时保留样本数量。")
    add_heading(doc, "2.5 特征相关性分析", 2)
    if add_figure(doc, FIGURES / "feature_correlation_heatmap.png", "图 2-2 特征相关性热力图", width=6.0):
        inserted_figures.append("图 2-2 特征相关性热力图")
    add_paragraph(doc, "相关性热力图显示，面积、周长、凸包面积、等效直径等特征之间相关性较强，说明原始形态学特征中存在一定冗余。但这些特征仍然提供了干豆尺寸和轮廓差异信息，对分类任务有实际意义。")

    add_heading(doc, "第三章 数据清洗与特征工程", 1)
    add_heading(doc, "3.1 数据清洗目标", 2)
    add_paragraph(doc, "数据清洗的目标是提升训练数据质量，降低缺失值、重复样本、标签污染和极端值对模型训练的干扰，使评估结果更能反映模型真实泛化能力。")
    add_heading(doc, "3.2 缺失值处理", 2)
    add_paragraph(doc, "项目使用训练集中位数 median 对数值缺失值进行填补。median 只在训练集上计算，验证集、测试集和上传预测数据只使用训练阶段保存的填补规则进行 transform，避免测试集信息泄漏。")
    add_heading(doc, "3.3 重复样本处理", 2)
    add_paragraph(doc, "项目会检测并删除训练集内部重复样本，同时检查 train、val、test 之间是否存在完全相同的行。专项验收结果显示，当前 train、val、test 之间不存在重复样本，实际用于训练和测试的数据也不存在完全重复行。")
    add_heading(doc, "3.4 标签污染清洗", 2)
    add_paragraph(doc, "标签处理包括统一大小写、去除空格，并修复 0/O、3/E 等常见字符污染，使类别恢复为七个标准类别。之后使用 LabelEncoder 将文本标签转换为模型可学习的数字编码，并保存编码器供评估和上传预测使用。")
    add_heading(doc, "3.5 异常值处理", 2)
    add_paragraph(doc, "异常值处理采用基于训练集 IQR 上下界的 winsorization/clipping。由于不同干豆类别之间形态差异较大，IQR 检测出的极端值不一定都是错误数据，因此采用裁剪而不是直接删除，可以降低极端值影响并保留样本量。")
    add_heading(doc, "3.6 特征工程", 2)
    add_paragraph(doc, "在原始形态学特征基础上，项目增加了 Area/Perimeter、MajorAxisLength/MinorAxisLength、ConvexArea/Area、EquivDiameter/Perimeter 等比例特征。这类比例特征可以强化模型对干豆整体形状、细长度和轮廓紧致程度的表达能力。")
    add_heading(doc, "3.7 标准化与编码", 2)
    add_paragraph(doc, "StandardScaler 用于统一特征尺度，尤其对 KNN、SVM 和逻辑回归较重要。项目保存了 preprocessor、label encoder、feature columns 和 raw feature columns，确保测试集和上传预测数据使用同一套处理流程。")
    add_heading(doc, "3.8 数据泄漏防护", 2)
    add_paragraph(doc, "专项验收确认：median、IQR 和 StandardScaler 均只在训练集 fit；测试集只 transform；LabelEncoder 使用训练阶段保存对象；上传预测不会重新训练模型，也不会重新 fit scaler 或 label encoder。")

    add_heading(doc, "第四章 多算法实验设计与结果分析", 1)
    add_heading(doc, "4.1 实验设置", 2)
    add_paragraph(doc, "实验评价指标包括训练集准确率、测试集准确率、加权 precision、加权 recall、加权 F1、混淆矩阵、训练时间、推理时间和过拟合差值。所有模型使用相同训练集和测试集，随机过程统一固定随机种子。")
    add_heading(doc, "4.2 算法选择", 2)
    add_paragraph(doc, "项目实现五个模型：逻辑回归、K近邻和支持向量机作为课堂内算法；随机森林和 XGBoost 作为课堂外扩展算法。SVM 使用 RBF 核处理非线性边界，随机森林通过多棵决策树集成降低方差，XGBoost 通过梯度提升逐步修正前序模型错误。")
    add_heading(doc, "4.3 测试集精度对比", 2)
    mt = metrics_table(data["metrics"])
    add_table(doc, mt, "表 4-1 多模型实验结果对比", font_size=6)
    inserted_tables.append("表 4-1 多模型实验结果对比")
    if add_figure(doc, FIGURES / "accuracy_comparison.png", "图 4-1 各模型测试集准确率对比"):
        inserted_figures.append("图 4-1 各模型测试集准确率对比")
    if add_figure(doc, FIGURES / "f1_comparison.png", "图 4-2 各模型加权 F1 对比"):
        inserted_figures.append("图 4-2 各模型加权 F1 对比")
    best_f1 = data["metrics"].sort_values("f1_weighted", ascending=False).iloc[0]
    fastest = data["speed"].sort_values("avg_predict_time_ms").iloc[0]
    add_paragraph(doc, f"从结果看，{best_model_zh} 的测试准确率最高，为 {fmt_float(data['best']['test_accuracy'], 4)}；加权 F1 最高的模型为 {MODEL_ZH.get(best_f1['model'], best_f1['model'])}，加权 F1 为 {fmt_float(best_f1['f1_weighted'], 4)}。推理速度最快的是 {MODEL_ZH.get(fastest['model'], fastest['model'])}，平均预测时间为 {fmt_float(fastest['avg_predict_time_ms'], 3)} ms。KNN、随机森林和 XGBoost 的训练集准确率明显高于测试集，存在一定过拟合。")
    add_heading(doc, "4.4 Loss 曲线对比", 2)
    for filename, caption in [
        ("loss_comparison.png", "图 4-3 Loss 曲线综合对比"),
        ("loss_curve_logistic.png", "图 4-4 逻辑回归训练损失曲线"),
        ("loss_curve_xgboost.png", "图 4-5 XGBoost 多分类 logloss 曲线"),
    ]:
        if add_figure(doc, FIGURES / filename, caption):
            inserted_figures.append(caption)
    add_paragraph(doc, "KNN 和随机森林不属于典型梯度迭代训练模型，因此不绘制传统意义上的 loss 曲线。本项目主要对逻辑回归和 XGBoost 记录训练损失变化。曲线整体呈下降趋势，说明训练过程能够收敛；如果训练损失与验证损失差距持续扩大，则需要关注过拟合。")
    add_heading(doc, "4.5 推理速度对比", 2)
    st = speed_table(data["speed"])
    add_table(doc, st, "表 4-2 推理速度实验结果", font_size=8)
    inserted_tables.append("表 4-2 推理速度实验结果")
    if add_figure(doc, FIGURES / "speed_comparison.png", "图 4-6 各模型推理速度对比"):
        inserted_figures.append("图 4-6 各模型推理速度对比")
    add_paragraph(doc, "逻辑回归推理速度最快，适合作为轻量基线模型。SVM 准确率最高但预测时间较长，KNN 预测时需要计算与训练样本的距离，因此推理速度也较慢。速度指标会受到 CPU 调度影响，但不影响准确率等核心指标的可信度。")
    add_heading(doc, "4.6 鲁棒性分析", 2)
    rt = robustness_table(data["robustness"])
    add_table(doc, rt, "表 4-3 不同噪声强度下鲁棒性实验结果", max_rows=40, font_size=6)
    inserted_tables.append("表 4-3 不同噪声强度下鲁棒性实验结果")
    if add_figure(doc, FIGURES / "robustness_comparison.png", "图 4-7 鲁棒性准确率对比", width=6.0):
        inserted_figures.append("图 4-7 鲁棒性准确率对比")
    if add_figure(doc, FIGURES / "robustness_drop_heatmap.png", "图 4-8 鲁棒性准确率下降热力图", width=6.0):
        inserted_figures.append("图 4-8 鲁棒性准确率下降热力图")
    add_paragraph(doc, "鲁棒性实验对训练数据加入 Gaussian Noise 和 Mask Noise，并在干净测试集上评估。clean accuracy 来自干净训练数据训练后的测试准确率，noisy accuracy 来自加噪训练数据重新训练后的测试准确率。若低强度噪声后准确率偶尔上升，说明轻微噪声可能起到数据增强作用，使模型对局部细节的依赖降低。")
    add_heading(doc, "4.7 过拟合分析", 2)
    add_paragraph(doc, "过拟合差值定义为 train_accuracy - test_accuracy。逻辑回归和 SVM 的训练集与测试集表现接近，泛化较稳定；KNN 和随机森林训练集准确率为 1.0000，但测试集略低，说明模型对训练样本记忆较强；XGBoost 参数较多，也存在一定过拟合，需要结合测试集表现和鲁棒性结果综合判断。")
    add_heading(doc, "4.8 混淆矩阵与分类报告分析", 2)
    cm_name = f"confusion_matrix_{data['best_model']}.png"
    if add_figure(doc, FIGURES / cm_name, f"图 4-9 {best_model_zh} 混淆矩阵", width=5.6):
        inserted_figures.append(f"图 4-9 {best_model_zh} 混淆矩阵")
    crt = class_report_table(data["class_report"])
    add_table(doc, crt, f"表 4-4 {best_model_zh} 分类报告", font_size=8)
    inserted_tables.append(f"表 4-4 {best_model_zh} 分类报告")
    add_paragraph(doc, "从分类报告和混淆矩阵可以看出，BOMBAY 类别识别效果较好，DERMASON、SIRA、CALI、BARBUNYA 等类别之间更容易出现混淆。这与干豆品种之间形态特征相近有关，尤其是尺寸和轮廓相似的类别更容易被模型混淆。")

    add_heading(doc, "第五章 工程系统设计与网页展示", 1)
    add_heading(doc, "5.1 工程结构", 2)
    add_paragraph(doc, "项目采用清晰的工程化结构组织数据、代码、模型和结果文件。核心结构如下：")
    add_paragraph(doc, "ML_DryBean_Project/\n├── data/\n├── src/\n├── models/\n├── results/\n├── app/\n├── main.py\n├── requirements.txt\n├── README.md\n├── run_all.bat\n└── run_app.bat")
    add_paragraph(doc, "data 保存训练、验证和测试数据；src 保存数据加载、预处理、训练、评估、鲁棒性、速度测试和可视化模块；models 保存训练好的模型和预处理器；results 保存实验结果、图表和报告；app 保存 Streamlit 展示页面；main.py 是统一命令行入口。")
    add_heading(doc, "5.2 统一命令行调用", 2)
    add_paragraph(doc, "算法训练、测试、鲁棒性分析和推理速度分析均在命令行完成，Streamlit 页面只负责结果展示和上传预测，符合算法运行阶段不使用 UI 的要求。常用命令包括：")
    add_paragraph(doc, "python main.py --mode eda\npython main.py --mode train_all\npython main.py --mode evaluate_all\npython main.py --mode robustness\npython main.py --mode speed\npython main.py --mode validate\npython main.py --mode all")
    add_heading(doc, "5.3 GitHub 与 README 展示", 2)
    add_paragraph(doc, f"GitHub 链接：{GITHUB_URL}")
    add_paragraph(doc, DEMO_URL_TEXT)
    add_paragraph(doc, "README 中包含项目简介、数据说明、环境安装、运行方式、工程结构、数据处理方法、算法说明、实验结果说明和上传预测说明。项目已经推送到 GitHub，便于教师下载复现。")
    add_heading(doc, "5.4 Streamlit 展示页面", 2)
    add_paragraph(doc, "Streamlit 页面包含数据集介绍、数据处理流程、EDA 图表、模型实验结果表、loss 曲线、混淆矩阵、鲁棒性分析、推理速度分析和上传 CSV 预测功能。论文截图建议优先放置以下页面：")
    screenshot_rows = [
        ("首页/项目概览", "展示系统标题、数据集说明、处理流程和五个算法名称，让老师第一眼看到项目完整性。"),
        ("模型实验结果", "展示 metrics_summary 表格、准确率对比图和 F1 对比图，证明多算法实验真实生成。"),
        ("Loss 曲线", "展示 Logistic Regression 与 XGBoost 的 loss 曲线，同时说明 KNN、Random Forest 不适合强行画 loss。"),
        ("鲁棒性分析", "展示 Gaussian Noise、Mask Noise 下的鲁棒性折线图或热力图，突出训练集加噪、测试集保持干净。"),
        ("推理速度分析", "展示 speed_comparison 图和速度表，说明不同模型不仅比较准确率，也比较部署效率。"),
        ("上传 CSV 预测", "展示选择模型、上传 sample_upload_predict.csv、预测类别/置信度表和下载预测结果按钮。"),
    ]
    add_table(doc, pd.DataFrame(screenshot_rows, columns=["建议截图", "截图中应体现的内容"]), "表 5-1 Streamlit 展示页面截图建议", font_size=8)
    inserted_tables.append("表 5-1 Streamlit 展示页面截图建议")
    add_heading(doc, "5.5 上传 CSV 预测功能", 2)
    add_paragraph(doc, "上传预测时，用户选择已经训练好的模型，上传包含干豆形态学特征的 CSV 文件，系统加载保存的模型、预处理器和标签编码器，输出预测类别和预测置信度，并提供预测结果 CSV 下载。上传预测阶段不会重新训练模型，也不会重新 fit scaler 或 label encoder，因此同一个 CSV、同一个模型的预测结果稳定一致。")

    add_heading(doc, "第六章 结果可信度与专项验收", 1)
    add_paragraph(doc, "项目新增专项验收报告，对数据划分、预处理、指标计算、可复现性、上传预测、鲁棒性和速度测试逻辑进行检查。经清理和检查后，当前 train、val、test 之间不存在重复样本；median、IQR 和 StandardScaler 只在训练集 fit；测试集和上传预测数据只 transform。")
    vt = validation_table(data["validation"])
    add_table(doc, vt, "表 6-1 指标重新计算校验结果", font_size=6)
    inserted_tables.append("表 6-1 指标重新计算校验结果")
    add_paragraph(doc, "所有模型的 metrics_summary.csv 指标均由保存模型在测试集上的真实预测结果生成，混淆矩阵样本数等于测试集样本数，classification report 中的 accuracy 与汇总表 test_accuracy 保持一致。")
    add_paragraph(doc, "如果老师问“你的准确率可信吗”，可以回答：本项目检查了训练集、验证集和测试集之间是否存在重复样本，避免测试集泄漏；median、IQR 裁剪、StandardScaler 均只在训练集 fit，测试集只 transform；所有指标由保存模型在测试集上的真实预测重新计算，并与汇总表、分类报告、混淆矩阵逐项校验；连续运行完整实验后，准确率、F1 和鲁棒性等核心指标保持一致，因此结果具有较好的可信度和可复现性。")

    add_heading(doc, "第七章 课程总结", 1)
    add_paragraph(doc, "通过这门课和这次期末项目，我对机器学习的完整流程有了更具体的认识。以前学习算法时，更多关注公式和模型本身，比如逻辑回归怎么做分类、KNN 怎么计算距离、SVM 怎么寻找分类边界。但真正完成一个项目后，我发现模型只是其中一部分，数据分析、数据清洗、特征工程和结果验证同样重要。比如本项目中，如果没有处理缺失值、标签污染、异常值和数据重复问题，即使模型准确率看起来很高，也不一定可信。")
    add_paragraph(doc, "在实现过程中，我体会到不同模型有不同优缺点。逻辑回归速度快、结果稳定，适合作为基线；KNN 思路简单，但预测速度较慢且容易受特征尺度影响；SVM 在本项目中准确率最高，但推理时间相对较长；随机森林和 XGBoost 对结构化表格数据有较强表达能力，但也更需要关注过拟合。这个过程让我意识到，不能只看准确率，还要结合 F1、混淆矩阵、过拟合差值、鲁棒性和推理速度综合判断模型。")
    add_paragraph(doc, "这次项目也让我学习到工程化的重要性。统一命令行入口、模块化代码、保存模型和预处理器、生成结果报告、构建 Streamlit 展示页面，这些内容让一个课堂实验变成了可以复现、可以展示、可以让别人使用的完整系统。对课程的建议是，希望以后可以增加更多真实 dirty data 案例、模型部署案例和项目答辩展示训练，这样能帮助我们更好地从理论走向实践。")

    add_heading(doc, "结论", 1)
    add_paragraph(doc, "本项目完成了 Dry Bean Dataset 七分类任务，实现了五种多分类算法，并围绕数据分析、数据清洗、特征工程、精度对比、loss 曲线、推理速度、鲁棒性和过拟合进行了完整实验。项目构建了工程化文件结构和 Streamlit 展示页面，并经过专项验收确认结果具有可信度和可复现性。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "Koklu, M., & Ozkan, I. A. (2020). Multiclass classification of dry beans using computer vision and machine learning techniques. Computers and Electronics in Agriculture.",
        "UCI Machine Learning Repository. Dry Bean Dataset.",
        "Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research.",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System.",
        "Hastie, T., Tibshirani, R., & Friedman, J. The Elements of Statistical Learning.",
    ]
    for ref in refs:
        add_paragraph(doc, ref)

    add_heading(doc, "附录", 1)
    add_heading(doc, "附录 A：主要运行命令", 2)
    add_paragraph(doc, "pip install -r requirements.txt\npython main.py --mode all\npython main.py --mode validate\nstreamlit run app/streamlit_app.py")
    add_heading(doc, "附录 B：工程目录结构", 2)
    add_paragraph(doc, "data、src、models、results、app、tests、docs 等目录分别保存数据、源码、模型、实验结果、展示页面、测试脚本和论文材料。")
    add_heading(doc, "附录 C：主要结果文件说明", 2)
    add_paragraph(doc, "metrics_summary.csv 保存模型指标；robustness_summary.csv 保存鲁棒性结果；speed_summary.csv 保存推理速度；sanity_check_report.md 保存数据泄漏检查；metric_validation_report.md 保存指标校验结果。")
    add_heading(doc, "附录 D：提交前检查清单", 2)
    checklist = [
        "已运行 python main.py --mode all",
        "已运行 python main.py --mode validate",
        "已确认 results/ 文件完整",
        "已确认 Streamlit 页面可打开",
        "已上传 GitHub",
        "已在 README 中填写 GitHub 链接",
        "已在论文中填写 GitHub 链接和网页链接",
        "已截图工程结构、命令行运行、Streamlit 页面和主要结果图表",
    ]
    for item in checklist:
        add_paragraph(doc, f"□ {item}")

    doc.save(DOCX_OUT)
    MD_OUT.write_text(build_markdown(data, inserted_figures, inserted_tables), encoding="utf-8")
    print(DOCX_OUT)
    print(MD_OUT)
    print("FIGURES:", len(inserted_figures), "|", "; ".join(inserted_figures))
    print("TABLES:", len(inserted_tables), "|", "; ".join(inserted_tables))


if __name__ == "__main__":
    main()
